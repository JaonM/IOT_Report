# -*- coding:utf-8 -*-
"""
report generalization service
"""
import datetime
import os
import re

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from app.configs.hotel_sensor_config import sensor_config
from app.connector.MySQLConnect import connect
from app.connector.MySQLConnect import connect2
from app.service.AnalyzeService import analyze_data_time_interval


def generalize_report(date=datetime.datetime.now().date().strftime('%Y-%m-%d'), total_result=None):
    if total_result is None:
        print('analyze result is not ready yet')
        return
    print('generating the report please wait...')

    dimensions = ['温度', '湿度', '电流', '电压', '功率', '电量', '时间频率']
    # main page
    for key in sensor_config.keys():
        document = Document()

        # 基地报告生成
        # add a cover page title
        document.add_paragraph()
        title = document.add_paragraph(key + ' 丽思卡尔顿传感器分析报告' + ' ' + date)
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_before = Pt(250)
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.name = 'Simhei'

        document.add_page_break()
        for result in total_result['results']:
            if result['company'] == key:
                # 基地传感器
                # if result['is_base'] == 1:
                document.add_heading(result['eui'], level=1)
                document.add_paragraph().paragraph_format.space_before = Pt(4)
                # try:
                #     document.add_picture(input_path + result['eui'] + '.png')
                # except FileNotFoundError:
                #     pass
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                if result['info'] == '缺失当天数据':
                    p = document.add_paragraph("缺失当天数据")
                    p.paragraph_format.left_indent = Inches(0.5)
                    continue
                # pics = os.listdir('../' + date)
                target_pic = None
                # for pic in pics:
                #     if re.search(result['eui'], pic):
                #         target_pic = pic
                #         break
                # if target_pic is not None:
                #     target_pic = '../' + date + '/' + target_pic
                #     p.alignment = WD_ALIGN_PARAGRAPH
                #     p.add_run().add_picture(target_pic, width=Inches(6.5))

                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.aligment = WD_ALIGN_PARAGRAPH

                for dim in dimensions:
                    try:
                        for msg in result['messages'][dim]:
                            if re.search('nan', msg['msg']):
                                continue
                            p.paragraph_format.left_indent = Inches(0.5)

                            run = p.add_run(msg['msg'])
                            if msg['status'] == 'red':  # DC143C
                                run.font.color.rgb = RGBColor(0xDC, 0x14, 0x3C)
                            elif msg['status'] == 'blue':
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
                    except KeyError:
                        continue
                document.add_page_break()
        # file_name = date.split()[0] + ' ' + key + '丽思卡尔顿传感器分析报告' + '.docx'
        document.save('./reports/' + date.split()[0] + '_' + key + '丽思卡尔顿传感器分析报告.docx')
    return date


def insert_report(file_name, file_path):
    """
    insert a report to database
    :param file_name:
    :return:
    """
    db = connect2()
    cursor = db.cursor()
    now = datetime.datetime.now()
    try:
        cursor.execute('insert into report (name,create_date,file_path) values(' + file_name + ',' + str(
            now) + ',' + file_path + ')')
        return 1
    except Exception:
        return -1


def create_data_img(date):
    """
    后台生成丢包缺失柱状图
    :param: date 日期
    :return:
    """
    try:
        img_list = []
        for key in sensor_config.keys():
            # for sensor in sensor_config[key]:
            #     eui = sensor['eui']
            if not os.path.exists('./reports/' + key + ' ' + date + '.csv'):
                analyze_data_time_interval(date=date)
            df = pd.read_csv('./reports/' + key + ' ' + date + '.csv', encoding='utf-8')

            for eui in df['eui'].unique():
                sub_df = df[df['eui'] == eui][['start_time', 'received_count']]
                sub_df.set_index(sub_df['start_time'], inplace=True)
                print(sub_df.head())
                sub_df.plot(kind='bar', title=str(eui) + ' ' + date)
                # plt.show()
                plt.tight_layout()
                plt.savefig('./app/static/data_img/' + str(eui) + ' ' + date + '.jpg')
                img_list.append('../static/data_img/' + str(eui) + ' ' + date + '.jpg')
        return img_list
    except Exception as e:
        print(e)
        return None


def load_data(start_date, end_date, page_num):
    """

    :param start_date:
    :param end_date:
    :param page_num: >1
    :return:
    """
    db = connect()
    cursor = db.cursor()

    if page_num == -1:
        cursor.execute("select tb2.eui,tb2.temperature 温度,tb2.humidity 湿度, tb2.batt 电量,tb2.current 电流,tb2.voltage 电压,tb2.power 功率,tb2.ts 获取时间,tb1.name 传感器名称,tb3.name 所属设备,tb4.name 所属酒店\
                                            from tb_sensor_data tb2 LEFT JOIN tb_sensor tb1 on tb1.code = tb2.eui LEFT JOIN tb_equipment tb3 on tb3.equipment_id = tb1.equipment_id \
                                            LEFT JOIN tb_customer tb4 on tb4.customer_id = tb1.customer_id where (tb2.eui='9896830000000008' or tb2.eui='9896830000000002' or tb2.eui = '9896830000000003' \
                                             or tb2.eui='9896830000000004' or tb2.eui = '9896830000000006' or tb2.eui = '3786E6ED0034004B' or \
                                             tb2.eui = '3768B26900230053' or tb2.eui = '4768B269002B0059' or tb2.eui = '4768B269001F003F' or \
                                             tb2.eui='4778B269003B002F' or tb2.eui='3430363057376506' or tb2.eui='3430363067378B07' or tb2.eui = '3430363064378607' \
                                             or tb2.eui='343036305D375E05' or tb2.eui = '3430363064378007') and tb2.ts > '" + start_date + "' and tb2.ts< '" + end_date + "'  ORDER BY tb2.ts ASC ")
    else:
        page_num -= 1
        cursor.execute("select tb2.eui,tb2.temperature 温度,tb2.humidity 湿度, tb2.batt 电量,tb2.current 电流,tb2.voltage 电压,tb2.power 功率,tb2.ts 获取时间,tb1.name 传感器名称,tb3.name 所属设备,tb4.name 所属酒店\
                                            from tb_sensor_data tb2 LEFT JOIN tb_sensor tb1 on tb1.code = tb2.eui LEFT JOIN tb_equipment tb3 on tb3.equipment_id = tb1.equipment_id \
                                            LEFT JOIN tb_customer tb4 on tb4.customer_id = tb1.customer_id where (tb2.eui='9896830000000008' or tb2.eui='9896830000000002' or tb2.eui = '9896830000000003' \
                                             or tb2.eui='9896830000000004' or tb2.eui = '9896830000000006' or tb2.eui = '3786E6ED0034004B' or \
                                             tb2.eui = '3768B26900230053' or tb2.eui = '4768B269002B0059' or tb2.eui = '4768B269001F003F' or \
                                             tb2.eui='4778B269003B002F' or tb2.eui='3430363057376506' or tb2.eui='3430363067378B07' or tb2.eui = '3430363064378607' \
                                             or tb2.eui='343036305D375E05' or tb2.eui = '3430363064378007') and tb2.ts > '" + start_date + "' and tb2.ts< '" + end_date + "'  ORDER BY tb2.ts ASC limit " + str(
            page_num) + ",20")
    result = list()
    for row in cursor.fetchall():
        result.append({'eui': row[0], 'temperature': row[1], 'humidity': row[2], 'battery': row[3], 'time': row[7]})
    return result


def alert(date):
    """
    丢包预警接口
    :param date: 日期 xxxx-xx-xx
    :return:
    """
    files = os.listdir('./app/static/data_img/')
    exist = False
    for file in files:
        if date in file:
            exist = True
            break
    if not exist:
        create_data_img(date)
    file_list = list()
    for file in files:
        if date in file:
            file_list.append(file)
    result_list = []
    for df in analyze_data_time_interval(date):
        for eui in df['eui'].unique():
            result = dict()
            result['eui'] = eui
            for file in file_list:
                if file.split()[0] == str(eui):
                    result['img_path'] = file
            sub_df = df[df['eui'] == eui]
            alert_list = list()
            for index, item in sub_df.iterrows():
                if 0.3 <= item['lost_rate'] < 0.6:
                    alert_list.append(
                        {'start_time': item['start_time'], 'end_time': item['end_time'], 'alert_info': '轻微丢包'})
                elif 0.6 <= item['lost_rate'] < 0.8:
                    alert_list.append(
                        {'start_time': item['start_time'], 'end_time': item['end_time'], 'alert_info': '中等严重丢包'})
                elif item['lost_rate'] >= 0.8:
                    alert_list.append(
                        {'start_time': item['start_time'], 'end_time': item['end_time'], 'alert_info': '严重丢包'})
                else:
                    alert_list.append(
                        {'start_time': item['start_time'], 'end_time': item['end_time'], 'alert_info': '正常接收'})
            result['alert'] = alert_list
            result_list.append(result)
    return result_list


if __name__ == '__main__':
    # create_data_img('2018-03-15')
    # print(load_data('2018-03-15', '2018-03-16', 1))
    print(alert('2018-03-15'))
